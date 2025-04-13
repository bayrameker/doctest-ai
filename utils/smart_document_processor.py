"""
NeuraParse Plus: Akıllı belge işleme modülü.
Bu modül, büyük belgeleri verimli bir şekilde işlemek için gelişmiş özellikler sunar:
- Verimli içerik çıkarma
- Akıllı özetleme
- Bellek-verimli akış işleme
- Yapılandırılmış veri çıkarma
- Çok büyük belgeleri bile analiz edebilme
"""

import logging
import os
import json
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple, Iterator
import time
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor

# Import document chunker
from utils.document_chunker import DocumentChunker, chunk_document_text

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants for document processing
TEMP_DIR = Path(tempfile.gettempdir()) / "document_processing"
SMART_PROCESSING_ENABLED = True  # Flag to enable/disable smart processing
STREAM_PROCESSING_THRESHOLD = 10 * 1024 * 1024  # 10MB threshold for stream processing
MAX_WORKERS = 4  # Max concurrent workers for parallel processing


class SmartDocumentProcessor:
    """
    NeuraParse Plus: Gelişmiş belge işleme motoru
    
    Akış, parçalama ve akıllı çıkarma teknikleriyle büyük belgeleri
    verimli bir şekilde işleyen gelişmiş belge işleme sınıfı.
    """
    
    def __init__(self, 
                 stream_threshold: int = STREAM_PROCESSING_THRESHOLD,
                 max_workers: int = MAX_WORKERS,
                 temp_dir: Optional[Path] = None):
        """
        Initialize the smart document processor.
        
        Args:
            stream_threshold: File size threshold for stream processing
            max_workers: Maximum number of concurrent workers
            temp_dir: Directory for temporary files
        """
        self.stream_threshold = stream_threshold
        self.max_workers = max_workers
        self.temp_dir = temp_dir or TEMP_DIR
        self.chunker = DocumentChunker()
        
        # Create temp directory if it doesn't exist
        os.makedirs(self.temp_dir, exist_ok=True)
        
        logger.info(f"NeuraParse Plus belge işleme motoru başlatıldı: akış_eşiği={stream_threshold}, paralel_işleyiciler={max_workers}")
    
    def should_stream_process(self, file_path: Union[str, Path]) -> bool:
        """
        Determine if a file should be processed in streaming mode.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            True if the file should be stream processed
        """
        file_size = os.path.getsize(file_path)
        return file_size > self.stream_threshold
    
    def extract_text_with_streaming(self, file_path: Union[str, Path]) -> Iterator[str]:
        """
        Extract text from a document using streaming for memory efficiency.
        
        Args:
            file_path: Path to the document file
            
        Yields:
            Text chunks from the document
        """
        try:
            # Determine file type
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()
            
            # Use appropriate parser based on file type
            if file_extension == '.pdf':
                yield from self._stream_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                yield from self._stream_docx(file_path)
            elif file_extension == '.txt':
                yield from self._stream_text(file_path)
            else:
                # Fall back to reading the whole file for unsupported formats
                logger.warning(f"Streaming not supported for {file_extension}, reading full file")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    yield f.read()
        except Exception as e:
            logger.error(f"Error in streaming text extraction: {str(e)}")
            yield f"Error extracting text: {str(e)}"
    
    def _stream_pdf(self, file_path: Path) -> Iterator[str]:
        """
        Stream text from a PDF file page by page.
        
        Args:
            file_path: Path to the PDF file
            
        Yields:
            Text from each page
        """
        try:
            # Import modules only when needed
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                for page_num in range(num_pages):
                    # Extract text from page
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        # Include page metadata
                        yield f"[Page {page_num + 1} of {num_pages}]\n{text}"
        except ImportError:
            logger.warning("PyPDF2 not available, falling back to standard processing")
            # Fall back to reading the file in chunks
            with open(file_path, 'rb') as file:
                while chunk := file.read(1024 * 1024):  # 1MB chunks
                    try:
                        yield chunk.decode('utf-8', errors='ignore')
                    except:
                        yield f"[Binary data chunk, decoding failed]"
    
    def _stream_docx(self, file_path: Path) -> Iterator[str]:
        """
        Stream text from a DOCX file paragraph by paragraph.
        
        Args:
            file_path: Path to the DOCX file
            
        Yields:
            Text from paragraphs
        """
        try:
            # Import modules only when needed
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            total_paragraphs = len(doc.paragraphs)
            current_section = ""
            section_count = 0
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Check if this looks like a heading
                if para.style.name.startswith('Heading'):
                    # Yield the previous section if it exists
                    if current_section:
                        yield current_section
                        current_section = ""
                    
                    # Start a new section with the heading
                    section_count += 1
                    current_section = f"[Section {section_count}] {text}\n\n"
                else:
                    # Add to current section
                    current_section += text + "\n\n"
                
                # Yield periodically to avoid building up too much text
                if i % 50 == 0 and current_section:
                    yield current_section
                    current_section = ""
            
            # Yield any remaining content
            if current_section:
                yield current_section
                
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_text = f"[Table {i+1}]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text += row_text + "\n"
                yield table_text
                
        except ImportError:
            logger.warning("python-docx not available, falling back to standard processing")
            # Fall back to reading the file in chunks
            try:
                import textract
                text = textract.process(str(file_path)).decode('utf-8')
                yield text
            except ImportError:
                logger.warning("textract not available, falling back to binary chunks")
                with open(file_path, 'rb') as file:
                    while chunk := file.read(1024 * 1024):  # 1MB chunks
                        try:
                            yield chunk.decode('utf-8', errors='ignore')
                        except:
                            yield f"[Binary data chunk, decoding failed]"
    
    def _stream_text(self, file_path: Path) -> Iterator[str]:
        """
        Stream a text file in chunks.
        
        Args:
            file_path: Path to the text file
            
        Yields:
            Text chunks
        """
        chunk_size = 1024 * 1024  # 1MB chunks
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            while chunk := file.read(chunk_size):
                yield chunk
    
    def process_document(self, file_path: Union[str, Path], extract_images: bool = False, 
                        extract_tables: bool = False, smart_chunking: bool = True) -> Dict[str, Any]:
        """
        Process a document with smart extraction and chunking.
        
        Args:
            file_path: Path to the document file
            extract_images: Whether to extract images
            extract_tables: Whether to extract tables
            smart_chunking: Whether to use smart semantic chunking
            
        Returns:
            Document processing results including text, structure, and extracted elements
        """
        start_time = time.time()
        file_path = Path(file_path)
        
        logger.info(f"NeuraParse Plus belge işleme başlatılıyor: {file_path} (görüntü_çıkarma={extract_images}, tablo_çıkarma={extract_tables})")
        
        # Initialize result structure
        result = {
            "filename": file_path.name,
            "file_type": file_path.suffix.lower().replace('.', ''),
            "file_size": os.path.getsize(file_path),
            "processed_at": time.time(),
            "processing_time": 0,
            "text": "",
            "chunks": [],
            "document_structure": {},
            "images": [],
            "tables": [],
            "metadata": {},
            "processing_method": "standard"
        }
        
        try:
            # Determine processing method based on file size
            if self.should_stream_process(file_path):
                logger.info(f"NeuraParse Plus akış işleme teknolojisi kullanılıyor - Büyük belge: {result['file_size'] / (1024*1024):.2f} MB")
                result["processing_method"] = "streaming"
                
                # Process with streaming
                text_chunks = []
                for chunk in self.extract_text_with_streaming(file_path):
                    text_chunks.append(chunk)
                
                result["text"] = "".join(text_chunks)
                
                # Apply smart chunking to the combined text if requested
                if smart_chunking and result["text"]:
                    result["chunks"] = self._apply_smart_chunking(result["text"])
                
            else:
                # Regular processing for smaller files
                result["processing_method"] = "standard"
                result.update(self._process_standard(file_path, extract_images, extract_tables))
                
                # Apply smart chunking if requested
                if smart_chunking and result["text"]:
                    result["chunks"] = self._apply_smart_chunking(result["text"])
            
            # Generate document structure map
            result["document_structure"] = self._extract_document_structure(result["text"])
            
            # Process images and tables in parallel if requested
            if extract_images or extract_tables:
                with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                    futures = []
                    
                    if extract_images:
                        futures.append(executor.submit(self._extract_images, file_path))
                    
                    if extract_tables:
                        futures.append(executor.submit(self._extract_tables, file_path))
                    
                    # Collect results
                    for future in concurrent.futures.as_completed(futures):
                        try:
                            data = future.result()
                            if "images" in data:
                                result["images"] = data["images"]
                            if "tables" in data:
                                result["tables"] = data["tables"]
                        except Exception as e:
                            logger.error(f"Error in parallel processing: {str(e)}")
            
            # Calculate processing time
            result["processing_time"] = time.time() - start_time
            
            logger.info(f"NeuraParse Plus belge işleme başarıyla tamamlandı - İşlem süresi: {result['processing_time']:.2f} saniye")
            return result
            
        except Exception as e:
            # Log error and return partial results with error info
            logger.error(f"NeuraParse Plus belge işleme hatası: {str(e)}")
            result["error"] = str(e)
            result["processing_time"] = time.time() - start_time
            return result
    
    def _process_standard(self, file_path: Path, extract_images: bool, extract_tables: bool) -> Dict[str, Any]:
        """
        Process document using standard (non-streaming) approach.
        
        Args:
            file_path: Path to the document file
            extract_images: Whether to extract images
            extract_tables: Whether to extract tables
            
        Returns:
            Processing results
        """
        result = {}
        
        # Determine file type
        file_extension = file_path.suffix.lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            result.update(self._process_pdf(file_path))
        elif file_extension in ['.docx', '.doc']:
            result.update(self._process_docx(file_path))
        elif file_extension == '.txt':
            # Simple text file
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                result["text"] = f.read()
        else:
            # Try generic extraction
            try:
                import textract
                result["text"] = textract.process(str(file_path)).decode('utf-8')
            except ImportError:
                logger.warning("textract not available, falling back to basic text extraction")
                with open(file_path, 'rb') as f:
                    result["text"] = f.read().decode('utf-8', errors='ignore')
        
        # Extract images if requested (but only if not already done in specific processors)
        if extract_images and "images" not in result:
            try:
                images_result = self._extract_images(file_path)
                result.update(images_result)
            except Exception as e:
                logger.error(f"NeuraParse Plus görüntü çıkarma hatası: {str(e)}")
                result["images"] = []
        
        # Extract tables if requested (but only if not already done in specific processors)
        if extract_tables and "tables" not in result:
            try:
                tables_result = self._extract_tables(file_path)
                result.update(tables_result)
            except Exception as e:
                logger.error(f"NeuraParse Plus tablo çıkarma hatası: {str(e)}")
                result["tables"] = []
        
        return result
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Process PDF document.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted data
        """
        result = {
            "text": "",
            "metadata": {},
            "page_count": 0
        }
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        if key.startswith('/'):
                            clean_key = key[1:]
                        else:
                            clean_key = key
                        
                        result["metadata"][clean_key] = str(value)
                
                # Extract page count
                result["page_count"] = len(reader.pages)
                
                # Extract text from all pages
                all_text = []
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    all_text.append(text)
                
                result["text"] = "\n\n".join(all_text)
                
            return result
            
        except ImportError:
            logger.warning("PyPDF2 not available, trying alternate PDF processing")
            
            try:
                # Try using textract as a fallback
                import textract
                result["text"] = textract.process(str(file_path)).decode('utf-8')
                return result
            except ImportError:
                logger.warning("textract not available, falling back to basic reading")
                with open(file_path, 'rb') as f:
                    result["text"] = f.read().decode('utf-8', errors='ignore')
                return result
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Process DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted data
        """
        result = {
            "text": "",
            "metadata": {},
            "sections": []
        }
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract core properties
            core_properties = ["author", "category", "comments", "content_status", 
                              "created", "identifier", "keywords", "language", 
                              "last_modified_by", "last_printed", "modified", 
                              "revision", "subject", "title", "version"]
            
            for prop in core_properties:
                try:
                    value = getattr(doc.core_properties, prop)
                    if value:
                        result["metadata"][prop] = str(value)
                except:
                    pass
            
            # Extract text from paragraphs
            text_parts = []
            current_section = {"heading": "", "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    continue
                
                # Add all text to the main text
                text_parts.append(text)
                
                # Track sections based on headings
                if para.style.name.startswith('Heading'):
                    # Save previous section if it exists and has content
                    if current_section["content"]:
                        result["sections"].append(current_section)
                    
                    # Start new section
                    current_section = {
                        "heading": text,
                        "level": int(para.style.name.replace("Heading", "")) if para.style.name.replace("Heading", "").isdigit() else 1,
                        "content": []
                    }
                else:
                    # Add to current section
                    current_section["content"].append(text)
            
            # Add final section if it has content
            if current_section["content"]:
                result["sections"].append(current_section)
            
            # Combine all text
            result["text"] = "\n\n".join(text_parts)
            
            return result
            
        except ImportError:
            logger.warning("python-docx not available, trying alternate DOCX processing")
            
            try:
                # Try using textract as a fallback
                import textract
                result["text"] = textract.process(str(file_path)).decode('utf-8')
                return result
            except ImportError:
                logger.warning("textract not available, falling back to basic reading")
                with open(file_path, 'rb') as f:
                    result["text"] = f.read().decode('utf-8', errors='ignore')
                return result
    
    def _extract_images(self, file_path: Path) -> Dict[str, List[Dict[str, Any]]]:
        """
        Extract images from document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with extracted images information
        """
        file_extension = file_path.suffix.lower()
        images = []
        
        try:
            if file_extension == '.pdf':
                images = self._extract_images_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                images = self._extract_images_from_docx(file_path)
            else:
                logger.warning(f"Image extraction not supported for {file_extension}")
        except Exception as e:
            logger.error(f"NeuraParse Plus görüntü çıkarma hatası: {str(e)}")
        
        return {"images": images}
    
    def _extract_images_from_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract images from PDF.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of extracted image information
        """
        images = []
        
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            for page_index in range(len(doc)):
                page = doc[page_index]
                
                # Get images
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    
                    if base_image:
                        # Save image to temporary file
                        image_extension = base_image["ext"]
                        image_filename = f"image_p{page_index+1}_{img_index+1}.{image_extension}"
                        temp_path = self.temp_dir / image_filename
                        
                        with open(temp_path, "wb") as f:
                            f.write(base_image["image"])
                        
                        # Get image dimensions
                        width = base_image.get("width", 0)
                        height = base_image.get("height", 0)
                        
                        # Add image info
                        images.append({
                            "filename": image_filename,
                            "path": str(temp_path),
                            "page": page_index + 1,
                            "index": img_index + 1,
                            "width": width,
                            "height": height,
                            "type": image_extension.upper(),
                            "size": len(base_image["image"])
                        })
            
            doc.close()
            
        except ImportError:
            logger.warning("PyMuPDF not available, trying alternate PDF image extraction")
            try:
                # Try using Pillow and PyPDF2 as fallback
                import PyPDF2
                from PIL import Image
                import io
                
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    for page_index in range(len(reader.pages)):
                        page = reader.pages[page_index]
                        
                        for img_index, resource in enumerate(page.get('/Resources', {}).get('/XObject', {}).values()):
                            if resource.get('/Subtype') == '/Image':
                                try:
                                    # Try to extract the image data
                                    data = resource.get_data()
                                    image_extension = 'png'  # Default
                                    
                                    # Save image to temporary file
                                    image_filename = f"image_p{page_index+1}_{img_index+1}.{image_extension}"
                                    temp_path = self.temp_dir / image_filename
                                    
                                    with open(temp_path, "wb") as f:
                                        f.write(data)
                                    
                                    # Add image info
                                    images.append({
                                        "filename": image_filename,
                                        "path": str(temp_path),
                                        "page": page_index + 1,
                                        "index": img_index + 1,
                                        "width": resource.get('/Width', 0),
                                        "height": resource.get('/Height', 0),
                                        "type": image_extension.upper(),
                                        "size": len(data)
                                    })
                                except:
                                    # Skip problematic images
                                    pass
            except ImportError:
                logger.warning("PyPDF2 and Pillow not available for image extraction")
        
        return images
    
    def _extract_images_from_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract images from DOCX.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of extracted image information
        """
        images = []
        
        try:
            from docx import Document
            import zipfile
            from PIL import Image
            import io
            
            # DOCX files are ZIP files containing images in word/media/
            temp_dir = self.temp_dir / "docx_images"
            os.makedirs(temp_dir, exist_ok=True)
            
            with zipfile.ZipFile(file_path) as docx_zip:
                # Find all image files in the zip
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for img_index, image_path in enumerate(image_files):
                    # Extract the image
                    image_data = docx_zip.read(image_path)
                    
                    # Get image filename from path
                    image_filename = os.path.basename(image_path)
                    temp_path = temp_dir / image_filename
                    
                    # Save image to temporary file
                    with open(temp_path, "wb") as f:
                        f.write(image_data)
                    
                    # Get image dimensions and type
                    try:
                        with Image.open(io.BytesIO(image_data)) as img:
                            width, height = img.size
                            img_format = img.format
                    except:
                        width, height = 0, 0
                        img_format = "UNKNOWN"
                    
                    # Add image info
                    images.append({
                        "filename": image_filename,
                        "path": str(temp_path),
                        "index": img_index + 1,
                        "width": width,
                        "height": height,
                        "type": img_format,
                        "size": len(image_data)
                    })
                    
        except ImportError:
            logger.warning("Required modules for DOCX image extraction not available")
        
        return images
    
    def _extract_tables(self, file_path: Path) -> Dict[str, List[Dict[str, Any]]]:
        """
        Extract tables from document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with extracted tables
        """
        file_extension = file_path.suffix.lower()
        tables = []
        
        try:
            if file_extension == '.pdf':
                tables = self._extract_tables_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                tables = self._extract_tables_from_docx(file_path)
            else:
                logger.warning(f"Table extraction not supported for {file_extension}")
        except Exception as e:
            logger.error(f"NeuraParse Plus tablo çıkarma hatası: {str(e)}")
        
        return {"tables": tables}
    
    def _extract_tables_from_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract tables from PDF.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of extracted tables
        """
        tables = []
        
        try:
            import camelot
            
            # Extract tables using camelot
            table_pages = camelot.read_pdf(str(file_path), pages='all', flavor='stream')
            
            for i, table in enumerate(table_pages):
                # Convert table to list format
                table_data = []
                df = table.df
                
                for _, row in df.iterrows():
                    table_data.append(list(row))
                
                tables.append({
                    "index": i + 1,
                    "page": table.page,
                    "data": table_data,
                    "shape": table.shape,
                    "accuracy": table.accuracy,
                    "whitespace": table.whitespace
                })
                
        except ImportError:
            logger.warning("camelot not available, trying alternate table extraction")
            try:
                # Try using tabula-py as fallback
                import tabula
                
                # Extract tables
                extracted_tables = tabula.read_pdf(str(file_path), pages='all')
                
                for i, df in enumerate(extracted_tables):
                    table_data = df.values.tolist()
                    
                    tables.append({
                        "index": i + 1,
                        "data": table_data,
                        "columns": df.columns.tolist(),
                        "shape": df.shape
                    })
                
            except ImportError:
                logger.warning("Neither camelot nor tabula are available for table extraction")
        
        return tables
    
    def _extract_tables_from_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Extract tables from DOCX.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of extracted tables
        """
        tables = []
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            for i, table in enumerate(doc.tables):
                # Convert table to data format
                table_data = []
                
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "index": i + 1,
                    "data": table_data,
                    "shape": (len(table.rows), len(table.rows[0].cells) if table.rows else 0)
                })
                
        except ImportError:
            logger.warning("python-docx not available for table extraction")
        
        return tables
    
    def _apply_smart_chunking(self, text: str) -> List[Dict[str, Any]]:
        """
        Apply smart chunking to document text.
        
        Args:
            text: Document text
            
        Returns:
            List of chunks with metadata
        """
        return chunk_document_text(text)
    
    def _extract_document_structure(self, text: str) -> Dict[str, Any]:
        """
        Extract document structure including headings, sections, etc.
        
        Args:
            text: Document text
            
        Returns:
            Document structure data
        """
        return DocumentChunker.generate_document_map(text)
    
    def save_processed_result(self, result: Dict[str, Any], output_dir: Optional[Path] = None) -> str:
        """
        Save processed document result to file.
        
        Args:
            result: Document processing result
            output_dir: Directory to save result (defaults to temp directory)
            
        Returns:
            Path to saved result file
        """
        # Use temp directory if output directory not specified
        save_dir = output_dir or self.temp_dir
        os.makedirs(save_dir, exist_ok=True)
        
        # Create filename based on original document name
        original_name = Path(result.get("filename", "document")).stem
        result_filename = f"{original_name}_processed_{int(time.time())}.json"
        output_path = save_dir / result_filename
        
        # Save result to file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, default=str)
        
        logger.info(f"Saved processing result to {output_path}")
        return str(output_path)


def smart_process_document(file_path: Union[str, Path], extract_images: bool = False,
                         extract_tables: bool = False) -> Dict[str, Any]:
    """
    Convenience function to process document with smart extraction.
    
    Args:
        file_path: Path to the document file
        extract_images: Whether to extract images
        extract_tables: Whether to extract tables
        
    Returns:
        Processing results
    """
    processor = SmartDocumentProcessor()
    return processor.process_document(file_path, extract_images, extract_tables)


def extract_document_structure(file_path: Union[str, Path]) -> Dict[str, Any]:
    """
    Extract document structure from a file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Document structure information
    """
    processor = SmartDocumentProcessor()
    result = processor.process_document(file_path, extract_images=False, extract_tables=False)
    return result["document_structure"]