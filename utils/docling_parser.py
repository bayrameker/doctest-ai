"""
Docling document parser modülü.
Bu modül, docling kütüphanesini veya basit doküman işleme kullanarak belgeleri LLM'ler için optimize edilmiş şekilde dönüştürür.
Sistem kaynakları yeterliyse gerçek Docling kullanır, değilse basit işleme yapar.
"""
import os
import logging
import json
import tempfile
import time
import threading
from typing import Dict, List, Any, Optional, Union, cast

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Yardımcı sınıflar - Tekrarlanan tanımlamaları önlemek için en üstte tanımlıyoruz
class DocumentWrapper:
    """Document sınıfı için temel wrapper"""
    def __init__(self, content):
        self.content = content
        self.items = []
        self.headings = []
        self.sections = []
        self.tables = []
        self.images = []
        
    def export_to_markdown(self):
        return self.content

class ResultWrapper:
    """Result sınıfı için temel wrapper"""
    def __init__(self, document):
        self.document = document

# Global variables for module availability
PDF_AVAILABLE = False
DOCX_AVAILABLE = False
PSUTIL_AVAILABLE = False
REAL_DOCLING_AVAILABLE = False

# Try to import PyPDF2 - koşullu import
try:
    import PyPDF2
    PDF_AVAILABLE = True
    logger.info("PyPDF2 modülü başarıyla yüklendi.")
except ImportError:
    PyPDF2 = None
    logger.warning("PyPDF2 modülü bulunamadı, PDF dosyaları işlenemeyecek")

# Try to import python-docx - koşullu import
try:
    import docx
    DOCX_AVAILABLE = True
    logger.info("python-docx modülü başarıyla yüklendi.")
except ImportError:
    docx = None
    logger.warning("python-docx modülü bulunamadı, DOCX dosyaları işlenemeyecek")

# Try to import psutil - koşullu import
try:
    import psutil
    PSUTIL_AVAILABLE = True
    logger.info("psutil modülü başarıyla yüklendi.")
except ImportError:
    psutil = None
    logger.warning("psutil modülü bulunamadı, bellek ve CPU kontrolü yapılamayacak")

# Try to import Docling
DocumentConverter = None
try:
    from docling.document_converter import DocumentConverter
    # Docling'in farklı sürümleri arasında uyum sağlamak için
    # Yeni importları koşullu olarak yapalım
    try:
        # Yeni Docling sürümlerinde bu modüller olabilir
        from docling.datamodel.pipeline_options import PdfPipelineOptions, TesseractCliOcrOptions
        logger.info("Docling pipeline_options modülleri başarıyla yüklendi.")
    except ImportError as e:
        logger.warning(f"Docling pipeline_options modülleri bulunamadı: {str(e)}")
        PdfPipelineOptions = None
        TesseractCliOcrOptions = None
        
    # Docling'in bazı sürümlerinde olmayabilir
    try:
        from docling.document_converter import PdfFormatOption
        logger.info("Docling PdfFormatOption başarıyla yüklendi.")
    except ImportError as e:
        logger.warning(f"Docling PdfFormatOption bulunamadı: {str(e)}")
        PdfFormatOption = None
        
    # Docling'in bazı sürümlerinde olmayabilir
    try:
        from docling.datamodel.base_models import InputFormat
        logger.info("Docling InputFormat başarıyla yüklendi.")
    except ImportError as e:
        logger.warning(f"Docling InputFormat bulunamadı: {str(e)}")
        InputFormat = None
        
    REAL_DOCLING_AVAILABLE = True
    logger.info("Gerçek Docling kütüphanesi başarıyla yüklendi.")
except ImportError as e:
    logger.warning(f"Gerçek Docling kütüphanesi yüklenemedi: {str(e)}")
    PdfPipelineOptions = None
    TesseractCliOcrOptions = None
    PdfFormatOption = None
    InputFormat = None

# Basit parser her zaman mevcuttur
LITE_DOCLING_AVAILABLE = PDF_AVAILABLE or DOCX_AVAILABLE

# Constants
DOCLING_TIMEOUT = 60  # Timeout in seconds - Daha uzun süre (1 dakika)
MIN_MEMORY_GB = 10  # Minimum memory in GB - Daha fazla bellek gereksinimi
MAX_DOCLING_FILE_SIZE = 5 * 1024 * 1024  # 5MB - Büyük dosyalar için doğrudan Lite mod
MAX_DOCLING_PAGES = 25  # Maximum number of pages to process with real Docling - Arttırıldı


def check_system_resources() -> bool:
    """
    Sistem kaynaklarını kontrol eder ve Docling için yeterli olup olmadığını değerlendirir
    
    Returns:
        bool: Kaynaklar yeterliyse True, değilse False
    """
    if not PSUTIL_AVAILABLE:
        logger.warning("psutil modülü bulunamadı, sistem kaynakları kontrol edilemiyor")
        return True  # Assume sufficient resources if psutil is not available
        
    try:
        # Check if psutil is available
        if psutil is None:
            logger.warning("psutil modülü kullanılamıyor, yeterli kaynak varsayılıyor")
            return True
            
        # Check available memory in GB
        if psutil is None:
            available_memory_gb = 32  # Varsayılan yüksek bellek değeri
        else:
            available_memory_gb = psutil.virtual_memory().available / (1024 * 1024 * 1024)
        
        # Check CPU usage - using a short interval to avoid blocking
        try:
            if psutil is None:
                cpu_usage = 0  # Varsayılan düşük CPU kullanımı
            else:
                cpu_usage = psutil.cpu_percent(interval=0.1)
        except Exception as cpu_err:
            logger.warning(f"CPU kullanımı kontrol edilemedi: {str(cpu_err)}")
            cpu_usage = 0  # Varsayılan değer
        
        # Evaluate both memory and CPU
        has_enough_memory = available_memory_gb >= MIN_MEMORY_GB
        has_low_cpu_load = cpu_usage < 80
        
        logger.info(f"Sistem kaynakları: {available_memory_gb:.2f} GB bellek müsait, CPU kullanımı: %{cpu_usage}")
        
        # CPU yüksek olsa bile bellek yeterliyse devam et
        if not has_low_cpu_load and has_enough_memory:
            logger.warning(f"CPU kullanımı yüksek (%{cpu_usage}) ancak yeterli bellek var, işleme devam ediliyor")
            return True
            
        # CPU yüksek ancak bellek yetersizse durumu hata olarak logla
        if not has_low_cpu_load and not has_enough_memory:
            logger.error(f"Sistem kaynakları yetersiz: CPU kullanımı %{cpu_usage}, bellek {available_memory_gb:.2f}GB (min: {MIN_MEMORY_GB}GB)")
        # Bellek yetersiz ancak CPU düşükse de durumu logla
        elif not has_enough_memory:
            logger.error(f"Yetersiz bellek: {available_memory_gb:.2f}GB (min: {MIN_MEMORY_GB}GB) - lite mod kullanılacak")
        # CPU yüksek, bellek yeterli ancak başka bir nedenden dolayı yeterli kaynak yoksa
        elif not has_low_cpu_load:
            logger.warning(f"CPU kullanımı çok yüksek: %{cpu_usage}")
            
        # Eğer bellek yetersiz ancak 5GB üzerindeyse, lite mod ile devam edebiliriz
        if not has_enough_memory and available_memory_gb >= 5:
            logger.warning(f"Bellek tam olarak yeterli değil ({available_memory_gb:.2f}GB) ancak lite mod için yeterli, devam edilecek")
            return True
            
        return has_enough_memory and has_low_cpu_load
    except Exception as e:
        logger.error(f"Sistem kaynaklarını kontrol ederken hata: {str(e)}")
        return False


def is_docling_available() -> bool:
    """
    Docling sisteminin (gerçek veya lite) kullanılabilir olup olmadığını kontrol eder
    
    Returns:
        bool: Herhangi bir Docling modu kullanılabilir ise True, değilse False
    """
    # At least lite mode should be available
    return LITE_DOCLING_AVAILABLE

# Docling'in hafif/lite modunun kullanılıp kullanılmadığını takip eden değişken
_using_lite_mode = False

def is_using_lite_mode() -> bool:
    """
    Docling'in hafif/lite modunun kullanılıp kullanılmadığını döndürür
    
    Returns:
        bool: Lite mod kullanıldıysa True, tam Docling kullanıldıysa False
    """
    return _using_lite_mode


def init_converter(pipeline_options=None, use_format_options=False):
    """
    DocumentConverter'ı güvenli bir şekilde başlat
    Tüm None kontrollerini yaparak, mevcut olmayan sınıfların çağrılmasını önler
    
    Args:
        pipeline_options: PdfPipelineOptions nesnesi
        use_format_options: Format options kullanımını zorlar
        
    Returns:
        Converter nesnesi veya None eğer converter oluşturulamazsa
    """
    # DocumentConverter sınıfı mevcut kontrol et
    if DocumentConverter is None:
        logger.error("DocumentConverter sınıfı bulunamadı, converter oluşturulamadı")
        return None
    
    # Converter oluşturmayı dene
    try:
        if pipeline_options is not None and not use_format_options:
            # Pipeline options ile converter oluştur
            try:
                converter = DocumentConverter(pipeline_options=pipeline_options)
                logger.info("DocumentConverter pipeline_options ile oluşturuldu")
                return converter
            except Exception as e:
                logger.warning(f"Pipeline options ile converter oluşturma hatası: {e}")
                # Düştü, format_options ile deneyeceğiz
                
        # Format options ile deneme
        if pipeline_options is not None and PdfFormatOption is not None and InputFormat is not None:
            try:
                format_options = {
                    InputFormat.PDF: PdfFormatOption(
                        pipeline_options=pipeline_options
                    )
                }
                # Format options ile converter oluştur
                converter = DocumentConverter(format_options=format_options)
                logger.info("DocumentConverter format_options ile oluşturuldu")
                return converter
            except Exception as e:
                logger.warning(f"Format options ile converter oluşturma hatası: {e}")
                # Düştü, varsayılan ile deneyeceğiz
                
        # Varsayılan converter oluştur
        try:
            converter = DocumentConverter()
            logger.info("DocumentConverter varsayılan ayarlarla oluşturuldu")
            return converter
        except Exception as e:
            logger.error(f"Varsayılan converter oluşturma hatası: {e}")
            return None
            
    except Exception as e:
        logger.error(f"Converter oluşturma ana hatası: {e}")
        return None


def use_real_docling(file_path: str = None) -> bool:
    """
    Gerçek Docling'in kullanılıp kullanılmayacağını belirler
    
    Args:
        file_path (str, optional): Belge dosyasının yolu, dosya boyutu kontrolü için
    
    Returns:
        bool: Gerçek Docling kullanılabilir ve kaynaklar yeterliyse True, değilse False
    """
    # Real Docling must be available
    if not REAL_DOCLING_AVAILABLE:
        logger.warning("Gerçek Docling mevcut değil")
        return False
    
    # Check system resources
    if not check_system_resources():
        logger.warning("Sistem kaynakları Docling için yeterli değil, lite mod kullanılacak")
        return False
    
    # Check file size if file_path is provided
    if file_path is not None and file_path:  # None ve boş string kontrolü
        try:
            file_size = os.path.getsize(file_path)
            if file_size > MAX_DOCLING_FILE_SIZE:
                logger.warning(f"Dosya boyutu ({file_size / (1024*1024):.2f} MB) çok büyük, Docling lite mod kullanılacak")
                return False
        except Exception as e:
            logger.error(f"Dosya boyutu kontrolünde hata: {str(e)}")
            # Continue with other checks if file size check fails
    
    # All conditions met
    return True


def parse_with_real_docling(file_path: str, max_tokens: int = 50000) -> Optional[Dict[str, Any]]:
    """
    Belgeyi gerçek Docling ile ayrıştırır
    
    Args:
        file_path (str): Belge dosyasının yolu
        max_tokens (int): İşlenecek maksimum token sayısı
        
    Returns:
        Optional[dict]: Belge içeriği ve yapılandırılmış veri veya hata durumunda None
    """
    # Gerçek mod kullanıldığını belirt
    global _using_lite_mode
    _using_lite_mode = False
    if not REAL_DOCLING_AVAILABLE or not DocumentConverter:
        logger.warning("Gerçek Docling kullanılamıyor")
        return None
        
    try:
        # Hugging Face model indirme kontrolü
        # Model kontrollü şekilde indir, önceden indirilip indirilmediğini kontrol et
        import os
        hf_home = os.environ.get("HF_HOME", os.path.expanduser("~/.cache/huggingface"))
        model_path = os.path.join(hf_home, "models--ds4sd--docling-models")
        
        if not os.path.exists(model_path):
            logger.warning("DocLing modelleri indirilmemiş, bu ilk kullanım için uzun sürebilir")
        
        # Dosya uzantısını kontrol et
        _, extension = os.path.splitext(file_path)
        extension = extension.lower().lstrip('.')
        
        logger.info(f"Belge {file_path} gerçek Docling ile ayrıştırılıyor")
        
        # Thread synchronization objects
        done_event = threading.Event()
        result_dict: Dict[str, Any] = {}
        errors: Dict[str, Optional[str]] = {"message": None}
        
        def docling_worker():
            try:
                # Öncelikle model yüklenmeden önce belgeleri incele
                file_size = os.path.getsize(file_path)
                if file_size > 10 * 1024 * 1024:  # 10MB'den büyük dosyalar için
                    logger.warning(f"Dosya boyutu çok büyük ({file_size / (1024*1024):.2f} MB), işlem uzun sürebilir")
                
                # Dosya tipini kontrol edelim
                _, extension = os.path.splitext(file_path)
                extension = extension.lower()
                
                # Dosya uzantısı geçerli mi?
                if extension not in ['.pdf', '.docx', '.doc', '.txt', '.md']:
                    logger.warning(f"Dosya tipi {extension} Docling tarafından desteklenmiyor olabilir")
                
                # Converter'ı PDF formatı için özel ayarlar ile başlat
                pipeline_options = None
                try:
                    if extension == '.pdf' and PdfPipelineOptions is not None:
                        # PDF dosyaları için yapılandırma ayarları
                        ocr_options = None
                        try:
                            # OCR için TesseractCliOcrOptions kullan
                            ocr_options = TesseractCliOcrOptions(lang=["auto"])
                            logger.info("TesseractCliOcrOptions başarıyla oluşturuldu")
                        except Exception as ocr_err:
                            logger.warning(f"TesseractCliOcrOptions oluşturma hatası: {str(ocr_err)}")
                        
                        # Pipeline options oluştur
                        try:
                            pipeline_options = PdfPipelineOptions(
                                do_ocr=True, 
                                force_full_page_ocr=True,
                                ocr_options=ocr_options,
                                generate_page_images=True
                            )
                            logger.info("PdfPipelineOptions başarıyla oluşturuldu")
                        except Exception as pipeline_err:
                            logger.warning(f"PdfPipelineOptions oluşturma hatası: {str(pipeline_err)}")
                        
                        # Format options oluştur - Tüm olası Docling sürümleri ile uyumlu
                        if pipeline_options is not None:
                            try:
                                # DocumentConverter sınıfı mevcut ve None değil kontrolü
                                if DocumentConverter is not None:
                                    # Sadece pipeline_options parametresi kullanmayı dene
                                    converter = DocumentConverter(pipeline_options=pipeline_options)
                                    logger.info("DocumentConverter pipeline_options ile başarıyla oluşturuldu")
                                else:
                                    logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                                    raise ValueError("DocumentConverter sınıfı bulunamadı")
                            except Exception as format_err1:
                                logger.warning(f"Pipeline options kullanım hatası: {str(format_err1)}")
                            
                            # PdfFormatOption ve InputFormat varsa format_options kullanmayı dene
                            if DocumentConverter is not None and PdfFormatOption is not None and InputFormat is not None:
                                try:
                                    format_options = {
                                        InputFormat.PDF: PdfFormatOption(
                                            pipeline_options=pipeline_options
                                        )
                                    }
                                    # Docling DocumentConverter'ı format options ile oluştur
                                    converter = DocumentConverter(format_options=format_options)
                                    logger.info("DocumentConverter format options ile başarıyla oluşturuldu")
                                except Exception as format_err2:
                                    logger.warning(f"Format options oluşturma hatası: {str(format_err2)}")
                                    # Fallback to default converter if available
                                    if DocumentConverter is not None:
                                        converter = DocumentConverter()
                                        logger.info("DocumentConverter varsayılan ayarlar ile oluşturuldu")
                                    else:
                                        logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                                        raise ValueError("DocumentConverter sınıfı bulunamadı")
                            else:
                                # Fallback to default converter if available
                                if DocumentConverter is not None:
                                    converter = DocumentConverter()
                                    logger.info("DocumentConverter varsayılan ayarlar ile oluşturuldu")
                                else:
                                    logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                                    raise ValueError("DocumentConverter sınıfı bulunamadı")
                        else:
                            # Fallback to default converter if available
                            if DocumentConverter is not None:
                                converter = DocumentConverter()
                                logger.info("DocumentConverter varsayılan ayarlar ile oluşturuldu")
                            else:
                                logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                                raise ValueError("DocumentConverter sınıfı bulunamadı")
                    else:
                        # Diğer dosya türleri için varsayılan belge dönüştürücü
                        if DocumentConverter is not None:
                            converter = DocumentConverter()
                            logger.info("DocumentConverter varsayılan ayarlar ile oluşturuldu (PDF olmayan dosya türü)")
                        else:
                            logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                            raise ValueError("DocumentConverter sınıfı bulunamadı")
                except Exception as converter_err:
                    logger.warning(f"DocumentConverter oluşturma hatası: {str(converter_err)}")
                    # Fallback to default converter if available
                    if DocumentConverter is not None:
                        converter = DocumentConverter()
                        logger.info("DocumentConverter varsayılan ayarlar ile oluşturuldu (hata telafisi)")
                    else:
                        logger.error("DocumentConverter sınıfı None, belge ayrıştırılamayacak")
                        raise ValueError("DocumentConverter sınıfı bulunamadı")
                
                # DOCX dosyası için python-docx ile içeriği yedekleme amaçlı çıkaralım
                backup_content = ""
                if extension == '.docx' and DOCX_AVAILABLE and docx is not None:
                    try:
                        doc = docx.Document(file_path)
                        for para in doc.paragraphs:
                            if para.text and para.text.strip():
                                backup_content += para.text.strip() + "\n\n"
                        logger.info(f"Yedek içerik çıkarıldı - {len(backup_content)} karakter")
                    except Exception as backup_err:
                        logger.warning(f"Yedek içerik çıkarma hatası: {str(backup_err)}")
                
                # Generator hatası için düzeltme
                try:
                    # DocumentConverter kullanımı farklı Docling sürümlerinde değişebilir
                    # Önce parametresiz kullanımı deneyin (sayfa sınırlaması olmadan)
                    logger.info(f"DocumentConverter.convert() çağrılıyor (sayfa sınırlaması olmadan)")
                    
                    # Geçerli sonuç kontrolü için bayrak
                    valid_result = False
                    
                    try:
                        # Docling 2.0+ versiyonu için max_pages parametresi deneme
                        result_obj = converter.convert(file_path, max_pages=MAX_DOCLING_PAGES)
                        valid_result = True
                    except TypeError:
                        # Eski Docling versiyonları için parametresiz kullanım
                        logger.info("max_pages parametresi desteklenmiyor, parametresiz çağrılıyor")
                        result_obj = converter.convert(file_path)
                        valid_result = True
                    except Exception as convert_err:
                        logger.error(f"convert() hatası: {str(convert_err)}")
                        # Geçersiz sonuç, yedek içeriği kullanacağız
                        valid_result = False
                    
                    # Docling 2.0+ için convert() davranışını kontrol et
                    # Dokümanın gösterdiği kullanım:
                    # result = converter.convert(source)
                    # print(result.document.export_to_markdown())
                    
                    # Olası durumlar:
                    # 1. Direkt olarak beklediğimiz formatta result döndürür ve result.document.export_to_markdown() çalışır
                    # 2. String döndürür (markdown veya düz metin)
                    # 3. Jeneratör döndürür ve list(result) kullanmak gerekir
                    # 4. Döküman üzerinden oluşturulmuş bir obje döndürür ve objeden içerik çıkarmak gerekir
                    
                    if isinstance(result_obj, str):
                        # Doğrudan metin döndürdüyse - basit durumda
                        logger.info("DocumentConverter string döndürdü, yapay document objesi oluşturuluyor")
                        markdown_content = result_obj  # Dönen stringi direkt kullan
                        
                        # Docling beklenen API'sine uygun wrapper oluştur
                        from collections import namedtuple
                        
                        # Yapay API oluştur - En üstte tanımladığımız DocumentWrapper ve ResultWrapper kullan
                        fake_document = DocumentWrapper(markdown_content)
                        result_obj = ResultWrapper(fake_document)
                        
                    elif hasattr(result_obj, '__iter__') and not hasattr(result_obj, 'document'):
                        # Jeneratör durumunu kontrol et
                        logger.info("DocumentConverter generator olabilir, kontrol ediliyor...")
                        
                        try:
                            # Jeneratörü güvenli bir şekilde listeye dönüştür
                            result_list = []
                            
                            try:
                                # Direkt listeye dönüştürmeye çalış
                                result_list = list(result_obj)
                            except Exception as list_err:
                                logger.warning(f"Generator direkt listeye dönüştürülemedi: {str(list_err)}")
                                
                                # Tek tek elemanları almayı dene
                                try:
                                    for item in result_obj:
                                        result_list.append(item)
                                except Exception as iter_err:
                                    logger.warning(f"Generator iterate edilemedi: {str(iter_err)}")
                            
                            if result_list and len(result_list) > 0:
                                logger.info(f"Generator {len(result_list)} eleman içeriyor")
                                
                                # İlk eleman analizi
                                first_item = result_list[0]
                                
                                if isinstance(first_item, str):
                                    # String elemanlar - birleştir
                                    logger.info("Generator string elemanlar içeriyor, birleştiriliyor")
                                    markdown_content = '\n\n'.join([str(item) for item in result_list if item])
                                    
                                    # Yapay API oluştur (String içerik için) - En üstte tanımladığımız sınıfları kullanarak
                                    fake_document = DocumentWrapper(markdown_content)
                                    result_obj = ResultWrapper(fake_document)
                                
                                elif hasattr(first_item, 'document'):
                                    # İlk eleman zaten document özelliğine sahip
                                    logger.info("Generator document özelliği olan eleman içeriyor")
                                    result_obj = first_item
                                    
                                else:
                                    # Bilinmeyen formatta sonuç
                                    logger.warning(f"Generator tanınmayan formatta elemanlar içeriyor: {type(first_item)}")
                                    
                                    # Tüm çıktıyı string olarak birleştirmeyi dene
                                    combined_text = '\n\n'.join([str(item) for item in result_list if item])
                                    
                                    # Yapay API oluştur (Bilinmeyen eleman formatı için) - En üstte tanımladığımız sınıfları kullanarak
                                    fake_document = DocumentWrapper(combined_text)
                                    result_obj = ResultWrapper(fake_document)
                            elif backup_content and len(backup_content) > 0:
                                # Boş liste ama yedek içerik var
                                logger.warning("Generator boş sonuç döndürdü, yedek içerik kullanılacak")
                                
                                # Yapay API oluştur (Yedek içerik için) - En üstte tanımladığımız sınıfları kullanarak
                                fake_document = DocumentWrapper(backup_content)
                                result_obj = ResultWrapper(fake_document)
                                
                                logger.info(f"Yedek içerik kullanıldı - {len(backup_content)} karakter")
                            else:
                                # Ne jeneratör ne de yedek içerik var
                                logger.warning("Generator boş sonuç döndürdü ve yedek içerik de bulunamadı")
                                raise ValueError("Docling sonuç listesi boş")
                                
                        except Exception as gen_error:
                            logger.error(f"Generator dönüşüm hatası: {str(gen_error)}")
                            logger.exception("Generator dönüşüm hatası detayları:")
                            
                            # Tüm Docling iç değişkenlerini logla
                            try:
                                logger.debug(f"Converter sınıfı: {type(converter)}")
                                logger.debug(f"Result objesi sınıfı: {type(result_obj)}")
                                
                                if hasattr(converter, 'pipeline_options'):
                                    logger.debug(f"Pipeline options: {converter.pipeline_options}")
                                
                                if hasattr(result_obj, 'document'):
                                    logger.debug(f"Result document sınıfı: {type(result_obj.document)}")
                            except Exception as debug_err:
                                logger.error(f"Hata değişkenlerini loglarken sorun: {str(debug_err)}")
                            
                            # Hatayla başa çıkmaya çalış - yedek içerik varsa onu kullan
                            content_to_use = ""
                            if backup_content and len(backup_content) > 100:
                                logger.info(f"Generator hatası oluştu, yedek içerik kullanılıyor - {len(backup_content)} karakter")
                                logger.debug(f"Yedek içerik başlangıcı: {backup_content[:200]}...")
                                content_to_use = backup_content
                            else:
                                content_to_use = f"Docling belge dönüşüm hatası: {str(gen_error)}"
                            
                            # En üstte tanımlanmış DocumentWrapper ve ResultWrapper sınıflarını kullanıyoruz
                            fake_document = DocumentWrapper(content_to_use)
                            result_obj = ResultWrapper(fake_document)
                except TypeError:
                    try:
                        # Eğer bu çalışmazsa başka alternatif parametrelerle deneyin
                        logger.info(f"DocumentConverter.convert() alternatif yöntemlerle deneniyor...")
                        
                        # Farklı parametreleri dene
                        try:
                            # Bazı Docling versiyonları için
                            result_obj = converter.convert(file_path, max_num_pages=MAX_DOCLING_PAGES)
                        except TypeError:
                            try:
                                # Başka bir alternatif
                                result_obj = converter.convert(file_path, pages=MAX_DOCLING_PAGES)
                            except TypeError:
                                # En son parametresiz dene
                                logger.info("Tüm parametreli yöntemler başarısız, parametresiz çağrılıyor")
                                result_obj = converter.convert(file_path)
                    except Exception as e:
                        logger.error(f"DocumentConverter.convert() alternatif yöntem hatası: {str(e)}")
                        # Yedek içerik varsa kullan
                        if backup_content and len(backup_content) > 100:
                            logger.info(f"Alternatif yöntem hatası, yedek içerik kullanılıyor - {len(backup_content)} karakter")
                            content_to_use = backup_content
                            
                            # En üstte tanımlanmış DocumentWrapper ve ResultWrapper sınıflarını kullanıyoruz
                            fake_document = DocumentWrapper(content_to_use)
                            result_obj = ResultWrapper(fake_document)
                            valid_result = True
                        else:
                            raise ValueError("Docling convert işlemi başarısız oldu")
                
                # result_obj tanımlı mı kontrol et
                if not result_obj:
                    logger.error("Docling sonucu oluşturulamadı")
                    raise ValueError("Docling sonucu oluşturulamadı")
                
                # document özelliği var mı kontrol et
                if not hasattr(result_obj, 'document'):
                    logger.error("Docling sonucu 'document' içermiyor")
                    raise ValueError("Docling sonucu beklenen formatta değil")
                    
                # Get markdown content
                markdown_content = result_obj.document.export_to_markdown()
                
                # If markdown content is empty, get raw text
                if not markdown_content.strip():
                    document = result_obj.document
                    text_content = ""
                    for item in document.items:
                        if hasattr(item, "text") and item.text:
                            text_content += item.text + "\n\n"
                    markdown_content = text_content if text_content.strip() else "Belge içeriği çıkarılamadı"
                
                # Calculate approximate token count
                token_count = len(markdown_content) // 4
                
                # Apply token limit
                if token_count > max_tokens:
                    markdown_content = markdown_content[:max_tokens * 4]
                    logger.info(f"İçerik {max_tokens} token sınırına göre kırpıldı")
                
                # Get document structure
                document = result_obj.document
                
                # Count sections and headings
                section_count = 0
                heading_count = 0
                
                if hasattr(document, "headings"):
                    heading_count = len(document.headings) if document.headings else 0
                    
                if hasattr(document, "sections"):
                    section_count = len(document.sections) if document.sections else 0
                
                # Count tables and images
                table_count = 0
                image_count = 0
                
                if hasattr(document, "tables"):
                    table_count = len(document.tables) if document.tables else 0
                    
                if hasattr(document, "images"):
                    image_count = len(document.images) if document.images else 0
                
                # Create document properties
                doc_result = {
                    "content": markdown_content,
                    "metadata": {
                        "filename": os.path.basename(file_path),
                        "file_type": extension,
                        "token_count": token_count,
                        "section_count": section_count,
                        "heading_count": heading_count,
                        "table_count": table_count,
                        "image_count": image_count
                    },
                    "elements": [],
                    "docling_parsed": True,
                    "docling_mode": "real"
                }
                
                # Create summary
                summary_parts = []
                summary_parts.append(f"Belge: {os.path.basename(file_path)}")
                
                if section_count > 0:
                    summary_parts.append(f"{section_count} bölüm")
                
                if heading_count > 0:
                    summary_parts.append(f"{heading_count} başlık")
                    
                if table_count > 0:
                    summary_parts.append(f"{table_count} tablo")
                    
                if image_count > 0:
                    summary_parts.append(f"{image_count} resim")
                    
                summary_parts.append(f"{token_count} token")
                
                doc_result["summary"] = ", ".join(summary_parts)
                
                # Update the shared result dict
                result_dict.update(doc_result)
                done_event.set()
            
            except Exception as e:
                errors["message"] = str(e)
                logger.error(f"Gerçek Docling hatası: {str(e)}")
                done_event.set()
        
        # Start the worker thread
        docling_thread = threading.Thread(target=docling_worker)
        docling_thread.daemon = True
        docling_thread.start()
        
        # Wait with timeout
        if not done_event.wait(timeout=DOCLING_TIMEOUT):
            logger.warning("Gerçek Docling zaman aşımına uğradı")
            return None
        
        # Check for errors
        if errors["message"] is not None:
            logger.warning(f"Gerçek Docling'de hata oluştu: {errors['message']}")
            return None
        
        # Check for empty result
        if not result_dict:
            logger.warning("Gerçek Docling boş sonuç döndürdü")
            return None
            
        return result_dict
    
    except Exception as e:
        logger.error(f"Gerçek Docling doküman işleme hatası: {str(e)}")
        return None


def parse_with_lite_docling(file_path: str, max_tokens: int = 50000) -> Dict[str, Any]:
    """
    Belgeyi basit yöntemlerle ayrıştırır ve yapılandırılmış veri olarak döndürür
    
    Args:
        file_path (str): Belge dosyasının yolu
        max_tokens (int): İşlenecek maksimum token sayısı (bellek kullanımını sınırlamak için)
        
    Returns:
        dict: Belge içeriği ve yapılandırılmış veri
    """
    # Lite mod kullanıldığını belirt
    global _using_lite_mode
    _using_lite_mode = True
    try:
        # Check file extension
        _, extension = os.path.splitext(file_path)
        extension = extension.lower().lstrip('.')
        
        logger.info(f"Belge {file_path} docling-lite ile ayrıştırılıyor")
        
        content = ""
        
        # For PDF files
        if extension == 'pdf' and PDF_AVAILABLE and PyPDF2 is not None:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() + "\n\n"
        
        # For DOCX files
        elif extension == 'docx' and DOCX_AVAILABLE and docx is not None:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        
        # For other file types
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            except Exception:
                # For binary files or encoding problems
                logger.warning(f"{file_path} dosyası metin formatında okunamadı")
                content = ""
        
        # Check if parsing failed
        if not content.strip():
            logger.warning("Dosya içeriği çıkarılamadı")
            return {
                "content": "",
                "metadata": {"error": "Dosya içeriği çıkarılamadı"},
                "elements": [],
                "docling_failed": True
            }
        
        # Calculate approximate token count
        token_count = len(content) // 4
        
        # Apply token limit
        if token_count > max_tokens:
            content = content[:max_tokens * 4]
            logger.info(f"İçerik {max_tokens} token sınırına göre kırpıldı")
        
        # Basic markdown conversion
        paragraphs = content.split('\n\n')
        markdown_content = content
        
        # Create artificial structure
        headings = []
        for i, paragraph in enumerate(paragraphs):
            if len(paragraph.strip()) > 0 and len(paragraph.strip()) < 100:
                if paragraph.isupper() or paragraph.endswith(':'):
                    headings.append(paragraph)
        
        # Create document properties
        result = {
            "content": markdown_content,
            "metadata": {
                "filename": os.path.basename(file_path),
                "file_type": extension,
                "token_count": token_count,
                "section_count": len(paragraphs),
                "heading_count": len(headings),
            },
            "elements": [],
            "docling_parsed": True,
            "docling_mode": "lite"
        }
        
        # Create summary
        summary_parts = []
        summary_parts.append(f"Belge: {os.path.basename(file_path)}")
        summary_parts.append(f"{len(paragraphs)} paragraf")
        
        if len(headings) > 0:
            summary_parts.append(f"{len(headings)} olası başlık")
            
        summary_parts.append(f"{token_count} token")
        
        result["summary"] = ", ".join(summary_parts)
        
        return result
    
    except Exception as e:
        logger.error(f"DoclingLite doküman işleme hatası: {str(e)}")
        return {
            "content": "",
            "metadata": {"error": str(e)},
            "elements": [],
            "docling_failed": True
        }


def parse_with_docling(file_path: str, max_tokens: int = 50000) -> Dict[str, Any]:
    """
    Belgeyi ayrıştırır ve yapılandırılmış veri olarak döndürür.
    Ortam şartlarına göre gerçek Docling veya lite mod arasında seçim yapar.
    
    Args:
        file_path (str): Belge dosyasının yolu
        max_tokens (int): İşlenecek maksimum token sayısı (bellek kullanımını sınırlamak için)
        
    Returns:
        dict: Belge içeriği ve yapılandırılmış veri
    """
    if not is_docling_available():
        logger.warning("Doküman işleme modülleri kullanılamıyor, temel bir yapı döndürülüyor")
        return {
            "content": "",
            "metadata": {"error": "Gerekli modüller bulunamadı veya yüklenemedi"},
            "elements": [],
            "docling_failed": True
        }
    
    # Determine which mode to use
    should_use_real_docling = use_real_docling(file_path)
    
    # If real Docling should be used
    if should_use_real_docling:
        logger.info("Sistem kaynakları yeterli, gerçek Docling kullanılıyor")
        
        # Try real Docling
        result = parse_with_real_docling(file_path, max_tokens)
        
        # Return if successful
        if result is not None:
            return result
        
        # Switch to lite mode if failed
        logger.warning("Gerçek Docling başarısız oldu, lite moda geçiliyor")
    else:
        logger.info("Sistem kaynakları yeterli değil, Docling lite modu kullanılıyor")
    
    # Use lite mode
    return parse_with_lite_docling(file_path, max_tokens)


def get_docling_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Belgenin yapısını analiz eder
    
    Args:
        file_path (str): Belge dosyasının yolu
        
    Returns:
        dict: Belge yapısı bilgisi
    """
    try:
        result = parse_with_docling(file_path)
        
        # Check if parsing failed
        if result.get("docling_failed", False):
            logger.warning("DoclingLite belge yapısı analizi başarısız oldu, temel yapı döndürülüyor")
            file_stats = os.stat(file_path)
            return {
                "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
                "content_size": file_stats.st_size,
                "is_llm_optimized": False,
                "docling_parse_error": result.get("metadata", {}).get("error", "Bilinmeyen hata")
            }
        
        # Extract document structure
        structure = {
            "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
            "content_size": len(result.get("content", "")),
            "is_llm_optimized": True,
            "parser": "docling_" + result.get("docling_mode", "lite")
        }
        
        # Get metadata
        metadata = result.get("metadata", {})
        
        # Add section count
        if "section_count" in metadata:
            structure["section_count"] = metadata["section_count"]
        
        # Add heading count
        if "heading_count" in metadata:
            structure["heading_count"] = metadata["heading_count"]
        
        # Add token count
        if "token_count" in metadata:
            structure["token_count"] = metadata["token_count"]
        
        return structure
    
    except Exception as e:
        logger.error(f"DoclingLite belge yapısı analizi hatası: {str(e)}")
        # Return basic structure on error
        return {
            "file_type": os.path.splitext(file_path)[1].lower().replace('.', ''),
            "content_size": os.path.getsize(file_path),
            "is_llm_optimized": False,
            "parser": "docling_lite",
            "docling_parse_error": str(e)
        }


def extract_docling_content(file_path: str) -> str:
    """
    Belgenin metin içeriğini çıkarır
    
    Args:
        file_path (str): Belge dosyasının yolu
        
    Returns:
        str: Belgenin metin içeriği
    """
    try:
        result = parse_with_docling(file_path)
        
        # Check if parsing failed
        if result.get("docling_failed", False):
            logger.warning("DoclingLite doküman içeriği çıkarılamadı, standart ayrıştırıcıya geçiliyor")
            return ""
        
        return result.get("content", "")
    
    except Exception as e:
        logger.error(f"DoclingLite doküman içeriği çıkarma hatası: {str(e)}")
        return ""